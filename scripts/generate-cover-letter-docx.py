#!/usr/bin/env python3
"""Extract cover letter from strategy markdown and render as ATS-friendly DOCX.

Parses Section 2 (Cover Letter Draft) from strategy document.
Calibri font, clean layout for ATS upload.

Usage: generate-cover-letter-docx.py <strategy.md>
"""

import re
import sys
import os
from datetime import date

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


# Colors
NAVY = RGBColor(30, 58, 95)
DARK_TEXT = RGBColor(15, 23, 42)
GRAY_TEXT = RGBColor(100, 116, 139)
BODY_TEXT = RGBColor(51, 65, 85)


def clean(text):
    """Strip markdown bold/italic markers."""
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"(?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*)", r"\1", text)
    return text.strip()


def extract_cover_letter(md_content):
    """Extract cover letter section from strategy markdown."""
    lines = md_content.split("\n")
    in_cover_letter = False
    in_subsection = None
    sections = {"opening": [], "body": [], "close": []}
    candidate_name = None

    for line in lines:
        stripped = line.strip()

        if re.match(r"^##\s+2\.\s+Cover Letter", stripped):
            in_cover_letter = True
            continue

        if in_cover_letter and re.match(r"^##\s+\d+\.", stripped):
            break

        if not in_cover_letter:
            continue

        if stripped.lower().startswith("### opening"):
            in_subsection = "opening"
            continue
        elif stripped.lower().startswith("### body"):
            in_subsection = "body"
            continue
        elif stripped.lower().startswith("### close"):
            in_subsection = "close"
            continue

        if stripped.startswith("Signed:"):
            continue

        if in_subsection and stripped:
            sections[in_subsection].append(clean(stripped))

    # Extract candidate name from close (last non-template line)
    if sections["close"]:
        for line in reversed(sections["close"]):
            if line and not any(word in line.lower() for word in ["thank", "look forward", "please", "best", "sincerely", "regards", "excited"]):
                candidate_name = line
                sections["close"].remove(line)
                break

    return sections, candidate_name


def main():
    if len(sys.argv) < 2:
        print("Usage: generate-cover-letter-docx.py <strategy.md>")
        sys.exit(1)

    md_path = sys.argv[1]
    if not os.path.exists(md_path):
        print(f"File not found: {md_path}")
        sys.exit(1)

    with open(md_path, "r") as f:
        md_content = f.read()

    sections, candidate_name = extract_cover_letter(md_content)

    if not any(sections.values()):
        print("Error: No cover letter section found in strategy document", file=sys.stderr)
        sys.exit(1)

    # Try to find resume in same directory for contact info
    strategy_dir = os.path.dirname(md_path)
    resume_files = [f for f in os.listdir(strategy_dir)
                    if f.endswith(".md") and "Strategy" not in f and "Changes" not in f and "CoverLetter" not in f]

    contact_line = None
    if resume_files:
        resume_path = os.path.join(strategy_dir, sorted(resume_files)[0])
        with open(resume_path) as f:
            for line in f:
                line = line.strip()
                if line.startswith("# "):
                    if not candidate_name:
                        candidate_name = clean(line[2:])
                elif "|" in line and ("@" in line or re.search(r"\d{3}", line)):
                    contact_line = clean(line)
                    break

    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(10)
    font.color.rgb = BODY_TEXT

    # Set margins
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Candidate name
    if candidate_name:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(candidate_name)
        run.font.size = Pt(16)
        run.font.bold = True
        run.font.color.rgb = DARK_TEXT
        run.font.name = "Calibri"
        p.paragraph_format.space_after = Pt(2)

    # Contact line
    if contact_line:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(contact_line)
        run.font.size = Pt(9)
        run.font.color.rgb = GRAY_TEXT
        run.font.name = "Calibri"
        p.paragraph_format.space_after = Pt(12)

    # Date
    today = date.today().strftime("%B %d, %Y")
    p = doc.add_paragraph()
    run = p.add_run(today)
    run.font.size = Pt(10)
    run.font.color.rgb = BODY_TEXT
    run.font.name = "Calibri"
    p.paragraph_format.space_after = Pt(12)

    # Salutation
    p = doc.add_paragraph()
    run = p.add_run("Dear Hiring Team,")
    run.font.size = Pt(10)
    run.font.color.rgb = DARK_TEXT
    run.font.name = "Calibri"
    p.paragraph_format.space_after = Pt(6)

    # Body paragraphs
    for section_name in ["opening", "body", "close"]:
        for para_text in sections[section_name]:
            p = doc.add_paragraph()
            run = p.add_run(para_text)
            run.font.size = Pt(10)
            run.font.color.rgb = BODY_TEXT
            run.font.name = "Calibri"
            p.paragraph_format.space_after = Pt(6)

    # Sign-off
    p = doc.add_paragraph()
    run = p.add_run("Sincerely,")
    run.font.size = Pt(10)
    run.font.color.rgb = DARK_TEXT
    run.font.name = "Calibri"
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)

    if candidate_name:
        p = doc.add_paragraph()
        run = p.add_run(candidate_name)
        run.font.size = Pt(10)
        run.font.bold = True
        run.font.color.rgb = DARK_TEXT
        run.font.name = "Calibri"

    # Output path
    base = os.path.basename(md_path)
    if "_Strategy_" in base:
        out_name = base.replace("_Strategy_", "_CoverLetter_").replace(".md", ".docx")
    else:
        out_name = base.replace("_application_strategy", "_CoverLetter").replace(".md", ".docx")

    docx_path = os.path.join(strategy_dir, out_name)
    doc.save(docx_path)
    print(docx_path)


if __name__ == "__main__":
    main()
