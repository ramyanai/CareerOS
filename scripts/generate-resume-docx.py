#!/usr/bin/env python3
"""Convert tailored resume markdown to an ATS-friendly DOCX.

Uses python-docx with Calibri font, standard heading styles, navy color scheme.
No headers/footers (ATS often ignores them).

Usage: generate-resume-docx.py <resume.md>
"""

import re
import sys
import os

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


# Colors
NAVY = RGBColor(30, 58, 95)
DARK_TEXT = RGBColor(15, 23, 42)
GRAY_TEXT = RGBColor(100, 116, 139)
BODY_TEXT = RGBColor(51, 65, 85)


def clean(text):
    """Strip markdown bold/italic markers."""
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"(?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*)", r"\1", text)
    return text.strip()


def parse_resume(md):
    """Parse resume markdown into structured sections."""
    lines = md.split("\n")
    sections = []
    current_section = None
    current_subsection = None
    name = None
    contact = None

    for line in lines:
        stripped = line.strip()

        if stripped.startswith("# ") and not stripped.startswith("## "):
            name = clean(stripped[2:])
            continue

        if name and not sections and not stripped.startswith("#") and stripped and (
            "|" in stripped or "@" in stripped or re.search(r"\d{3}[-.]?\d{3}[-.]?\d{4}", stripped)
        ):
            contact = clean(stripped)
            continue

        if stripped.startswith("## "):
            current_section = {"title": clean(stripped[3:]), "subsections": [], "bullets": [], "text": []}
            current_subsection = None
            sections.append(current_section)
            continue

        if stripped.startswith("### ") and current_section is not None:
            current_subsection = {"title": clean(stripped[4:]), "bullets": [], "text": []}
            current_section["subsections"].append(current_subsection)
            continue

        if (stripped.startswith("- ") or stripped.startswith("* ")) and current_section is not None:
            bullet_text = clean(stripped[2:])
            if current_subsection is not None:
                current_subsection["bullets"].append(bullet_text)
            else:
                current_section["bullets"].append(bullet_text)
            continue

        if stripped and current_section is not None:
            if current_subsection is not None:
                current_subsection["text"].append(clean(stripped))
            else:
                current_section["text"].append(clean(stripped))

    return name, contact, sections


def add_section_border(paragraph):
    """Add a bottom border to a paragraph (section heading underline)."""
    from docx.oxml.ns import qn
    pPr = paragraph._element.get_or_add_pPr()
    pBdr = pPr.makeelement(qn("w:pBdr"), {})
    bottom = pBdr.makeelement(qn("w:bottom"), {
        qn("w:val"): "single",
        qn("w:sz"): "4",
        qn("w:space"): "1",
        qn("w:color"): "1E3A5F",
    })
    pBdr.append(bottom)
    pPr.append(pBdr)


def main():
    if len(sys.argv) < 2:
        print("Usage: generate-resume-docx.py <resume.md>")
        sys.exit(1)

    md_path = sys.argv[1]
    if not os.path.exists(md_path):
        print(f"File not found: {md_path}")
        sys.exit(1)

    with open(md_path) as f:
        md = f.read()

    name, contact, sections = parse_resume(md)

    doc = Document()

    # Set default font to Calibri
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(10)
    font.color.rgb = BODY_TEXT

    # Set narrow margins
    for section in doc.sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    # Name
    if name:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(name)
        run.font.size = Pt(16)
        run.font.bold = True
        run.font.color.rgb = DARK_TEXT
        run.font.name = "Calibri"
        p.paragraph_format.space_after = Pt(2)

    # Contact line
    if contact:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(contact)
        run.font.size = Pt(9)
        run.font.color.rgb = GRAY_TEXT
        run.font.name = "Calibri"
        p.paragraph_format.space_after = Pt(6)

    # Sections
    for sec in sections:
        # Section heading
        p = doc.add_paragraph()
        run = p.add_run(sec["title"])
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = NAVY
        run.font.name = "Calibri"
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(2)
        add_section_border(p)

        # Section text
        for text in sec["text"]:
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.font.size = Pt(10)
            run.font.color.rgb = BODY_TEXT
            run.font.name = "Calibri"
            p.paragraph_format.space_after = Pt(2)

        # Section bullets
        for bullet in sec["bullets"]:
            p = doc.add_paragraph(style="List Bullet")
            p.clear()
            run = p.add_run(bullet)
            run.font.size = Pt(10)
            run.font.color.rgb = BODY_TEXT
            run.font.name = "Calibri"
            p.paragraph_format.space_after = Pt(1)

        # Subsections
        for sub in sec["subsections"]:
            p = doc.add_paragraph()
            run = p.add_run(sub["title"])
            run.font.size = Pt(10.5)
            run.font.bold = True
            run.font.color.rgb = DARK_TEXT
            run.font.name = "Calibri"
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(1)

            for text in sub["text"]:
                p = doc.add_paragraph()
                run = p.add_run(text)
                run.font.size = Pt(9)
                run.font.italic = True
                run.font.color.rgb = GRAY_TEXT
                run.font.name = "Calibri"
                p.paragraph_format.space_after = Pt(1)

            for bullet in sub["bullets"]:
                p = doc.add_paragraph(style="List Bullet")
                p.clear()
                run = p.add_run(bullet)
                run.font.size = Pt(10)
                run.font.color.rgb = BODY_TEXT
                run.font.name = "Calibri"
                p.paragraph_format.space_after = Pt(1)

    # Output
    docx_path = md_path.replace(".md", ".docx")
    doc.save(docx_path)
    print(docx_path)


if __name__ == "__main__":
    main()
